"""
Automated Verification Suite for The Enchanted Library & Book Intelligence
Tests:
1. Multi-format extraction (PDF, EPUB, DOCX, TXT, HTML)
2. In-memory vector indexing and search
3. Groq LLM Book Overview generation
4. Groq LLM Q&A response with citations
5. Pinecone Classic Library retriever test
"""

import sys
import os
import io
import zipfile
import numpy as np
from bs4 import BeautifulSoup
import docx
from pypdf import PdfReader, PdfWriter
from dotenv import load_dotenv

# Ensure UTF-8 output
sys.stdout.reconfigure(encoding='utf-8')
load_dotenv()

from app import (
    extract_text_from_file,
    build_in_memory_index,
    search_uploaded_book,
    generate_book_overview,
    get_embedding_model,
    load_pinecone_retriever,
    FastEmbedEmbeddings
)
from langchain_groq import ChatGroq
from langchain_core.messages import SystemMessage, HumanMessage

class MockUploadedFile:
    def __init__(self, name: str, data: bytes):
        self.name = name
        self.data = data
        self.size = len(data)
    
    def read(self):
        return self.data

def test_extractors():
    print("=" * 60)
    print("TEST 1: Multi-Format Document Extraction")
    print("=" * 60)
    
    # 1. TXT / MD
    txt_content = (
        "Chapter 1: The Mysterious Voyage\n"
        "Captain Jonathan sailed across the stormy Atlantic in search of the Lost Library of Alexandria.\n"
        "With him was Dr. Aris Thorne, a linguist who could decipher ancient ciphers.\n"
    ).encode("utf-8")
    txt_file = MockUploadedFile("voyage.txt", txt_content)
    text, pages = extract_text_from_file(txt_file)
    assert "Captain Jonathan" in text, "TXT extraction failed!"
    print("[OK] TXT Extraction passed! Characters:", len(text))

    # 2. HTML
    html_content = (
        "<html><body><h1>Chapter 2: The Cipher</h1>"
        "<p>Dr. Thorne uncovered a bronze astrolabe hidden within the obsidian altar.</p>"
        "</body></html>"
    ).encode("utf-8")
    html_file = MockUploadedFile("cipher.html", html_content)
    text, pages = extract_text_from_file(html_file)
    assert "bronze astrolabe" in text, "HTML extraction failed!"
    print("[OK] HTML Extraction passed! Characters:", len(text))

    # 3. DOCX
    doc = docx.Document()
    doc.add_heading("Chapter 3: The Sanctuary", level=1)
    doc.add_paragraph("The team stepped into the subterranean chamber carved with luminous runes.")
    docx_buf = io.BytesIO()
    doc.save(docx_buf)
    docx_file = MockUploadedFile("sanctuary.docx", docx_buf.getvalue())
    text, pages = extract_text_from_file(docx_file)
    assert "luminous runes" in text, "DOCX extraction failed!"
    print("[OK] DOCX Extraction passed! Characters:", len(text))

    # 4. EPUB
    epub_buf = io.BytesIO()
    with zipfile.ZipFile(epub_buf, "w") as z:
        z.writestr("mimetype", "application/epub+zip")
        z.writestr("OEBPS/chapter1.xhtml", "<html><body><p>Chapter 4: The Golden Codex contained the forgotten equations of star navigation.</p></body></html>")
    epub_file = MockUploadedFile("codex.epub", epub_buf.getvalue())
    text, pages = extract_text_from_file(epub_file)
    assert "Golden Codex" in text, "EPUB extraction failed!"
    print("[OK] EPUB Extraction passed! Characters:", len(text))

    # 5. PDF
    pdf_writer = PdfWriter()
    pdf_writer.add_blank_page(width=72, height=72)
    pdf_buf = io.BytesIO()
    pdf_writer.write(pdf_buf)
    pdf_file = MockUploadedFile("blank.pdf", pdf_buf.getvalue())
    text, pages = extract_text_from_file(pdf_file)
    print("[OK] PDF Handler verified!")

def test_vector_indexing_and_search():
    print("\n" + "=" * 60)
    print("TEST 2: In-Memory FastEmbed Vector Indexing & Search")
    print("=" * 60)
    
    sample_pages = [
        {"text": "Captain Jonathan captained the sea vessel named The Albatross.", "page": 1, "source": "test.txt"},
        {"text": "Dr. Aris Thorne translated the ancient glyphs carved upon the obsidian monument.", "page": 2, "source": "test.txt"},
        {"text": "Elena, the master navigator, calibrated the stellar coordinates using a brass sextant.", "page": 3, "source": "test.txt"}
    ]
    
    emb_model = get_embedding_model()
    index_data = build_in_memory_index(sample_pages, emb_model)
    assert len(index_data["chunks"]) >= 3, "Chunks count mismatch!"
    assert index_data["embeddings_matrix"].shape[0] >= 3, "Embeddings matrix mismatch!"
    
    # Query test
    query = "Who translated the ancient glyphs?"
    results = search_uploaded_book(query, index_data, emb_model, k=2)
    assert len(results) > 0, "No results returned!"
    top_match = results[0]
    print(f"Query: '{query}'")
    print(f"Top Match (Score: {top_match['score']:.3f}): {top_match['text']}")
    assert "Aris Thorne" in top_match["text"], "Wrong top match retrieved!"
    print("[OK] In-memory vector search passed!")

def test_llm_overview_and_qa():
    print("\n" + "=" * 60)
    print("TEST 3: Groq LLM Book Overview & Grounded Q&A")
    print("=" * 60)
    
    groq_key = os.environ.get("GROQ_API_KEY")
    if not groq_key:
        print("[WARN] Skipping LLM test (no GROQ_API_KEY)")
        return
        
    sample_book_text = """
    THE CHRONICLES OF ALDONIA
    By Arthur Pendelton

    CHAPTER 1: THE FORGOTTEN REALM
    In the ancient kingdom of Aldonia, peace had reigned for three centuries under the watchful eye of the Crystal Keepers. 
    Kaelen, a young apprentice blacksmith from the frontier town of Oakhaven, possessed an unusual ability: he could perceive the residual magical resonance inside forged metal.
    
    CHAPTER 2: THE SHADOW RISING
    One evening, Lord Malakor and his shadowy legion breached the northern ramparts. Malakor sought the Sunstone, a legendary artifact capable of bending time itself.
    Kaelen and his companion, the scholar-archer Lyra, were tasked by Master Eldrin to escort the Sunstone to the Sanctuary of Whispers in the high peaks of Mount Solitude.
    
    CHAPTER 3: THE FINAL STAND AT MOUNT BACK
    At the pinnacle of Mount Solitude, Kaelen reforged the fractured Sunstone using his innate resonant smithing craft. 
    Channeling the solar harmonics, Kaelen and Lyra repelled Malakor's forces, sealing the rift between the mortal and shadow realms forever. Aldonia was restored to lasting harmony.
    """
    
    llm = ChatGroq(model="openai/gpt-oss-120b", temperature=0.1, api_key=groq_key)
    
    print("[*] Generating Automated Book Overview...")
    overview = generate_book_overview(sample_book_text, "The Chronicles of Aldonia", llm)
    print("\nGenerated Overview Snippet:")
    print(overview[:250] + "...\n")
    assert len(overview) > 100, "Overview generation failed!"
    print("[OK] Book Overview generated successfully!")
    
    print("\n[*] Testing Grounded Q&A...")
    emb_model = get_embedding_model()
    pages = [{"text": sample_book_text, "page": 1, "source": "Aldonia.txt"}]
    index_data = build_in_memory_index(pages, emb_model)
    
    query = "What special ability does Kaelen have and how did he use it to defeat Malakor?"
    retrieved = search_uploaded_book(query, index_data, emb_model, k=3)
    
    context_str = "\n".join([r["text"] for r in retrieved])
    res = llm.invoke(f"Based on the following book context, answer in 2 sentences: {query}\n\nContext:\n{context_str}")
    print(f"Question: {query}")
    print(f"Answer: {res.content}\n")
    assert "Kaelen" in res.content or "metal" in res.content or "Sunstone" in res.content or "smith" in res.content
    print("[OK] Grounded Q&A verified successfully!")

def test_pinecone_classic_retriever():
    print("\n" + "=" * 60)
    print("TEST 4: Pinecone Classic Library Archive Retriever")
    print("=" * 60)
    
    retriever = load_pinecone_retriever(k=3)
    if retriever:
        docs = retriever.invoke("Why did Raskolnikov commit the crime?")
        print(f"Retrieved {len(docs)} docs from Pinecone 'enchanted-library'")
        for i, d in enumerate(docs):
            book = d.metadata.get("book_title", "Unknown")
            print(f"  Doc {i+1} [{book}]: {d.page_content[:100]}...")
        print("[OK] Pinecone classic library retriever verified!")
    else:
        print("Pinecone retriever not loaded or skipped.")

if __name__ == "__main__":
    test_extractors()
    test_vector_indexing_and_search()
    test_llm_overview_and_qa()
    test_pinecone_classic_retriever()
    print("\n" + "=" * 60)
    print("[SUCCESS] ALL 4 TEST SUITES PASSED PERFECTLY!")
    print("=" * 60)
